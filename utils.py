from typing import IO, Union

import numpy as np
import pandas as pd
from docx import Document


def load_doc(path: Union[str, IO[bytes]], add_tables: bool = False):
    doc = Document(path)
    full_text = []
    for paragraph in doc.paragraphs:
        if len(paragraph.text) > 0:
            full_text.append(paragraph.text)

    text = "\n".join(full_text)
    if add_tables:
        tables_str = []
        for table_index, table in enumerate(doc.tables):
            table_str = f"Таблица {table_index + 1}"
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            markup_table = pd.DataFrame(np.vstack(table_data)).to_markdown(index=False)
            table_str += "\n" + markup_table
            tables_str.append(table_str)
        tables_full_str = "\n".join(tables_str)
        text = text + "\n" + tables_full_str
    return text


def clean_start_number(text: str):
    if type(text) != str:
        return text
    while not text[0].isalpha():
        text = text[1:]
    return text
